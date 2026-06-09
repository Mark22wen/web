"""
knowledge_ingest.py — 把 public/资料/ 下的 PDF / Word / Excel 文档写入 ChromaDB

用法：
    python knowledge_ingest.py                  # 增量入库（已有 chunk 跳过）
    python knowledge_ingest.py --force          # 清空旧知识文档后全量重建
    python knowledge_ingest.py --update-meta    # 仅补写新 metadata 字段，不重新 embedding（快）
    python knowledge_ingest.py --file 报告.pdf  # 只处理指定文件（部分名匹配）
    python knowledge_ingest.py --no-vision      # 跳过 llava 图表识别

依赖：
    pip install pdfplumber pymupdf python-docx requests openpyxl

Metadata 字段（每个 chunk）：
    table         : "knowledge"（固定，用于与结构化数据隔离）
    source        : 相对路径
    filename      : 文件名
    chunk         : chunk 序号
    doc_year      : 文档发布年份（从文件名或首段内容提取）
    page          : 印刷页码（仅 PDF）
    section       : 当前章节标题（PDF / Word 均追踪）
    countries     : chunk 中提到的国家/地区（逗号分隔，最多6个）
    year_in_chunk : chunk 文本中出现最多的年份
"""

import os
import re
import time
import hashlib
import base64
import requests
from collections import Counter

# ── 配置 ──────────────────────────────────────────────
DOCS_FOLDER   = os.path.join(os.path.dirname(__file__), "public", "资料")
CHROMA_HOST   = os.environ.get("CHROMA_HOST", "localhost")
CHROMA_PORT   = int(os.environ.get("CHROMA_PORT", "8000"))
OLLAMA_URL    = os.environ.get("OLLAMA_URL", "http://localhost:11434")
EMBED_MODEL   = os.environ.get("OLLAMA_EMBED_MODEL", "nomic-embed-text")
COLLECTION    = "patent_knowledge"
TENANT        = "default_tenant"
DATABASE      = "default_database"
CHUNK_SIZE    = 600
CHUNK_OVERLAP = 100
BATCH_SIZE    = 10
EMBED_TIMEOUT = 30
use_no_vision = False   # --no-vision 标志，main() 中覆盖


# ── 元数据提取辅助 ────────────────────────────────────
# 与 server.js GLOBAL_COUNTRIES_RE 同步，用于 chunk 国家标注
COUNTRY_KEYWORDS = [
    '德国', '美国', '日本', '欧洲', '英国', '法国', '韩国', '印度', '俄罗斯',
    '意大利', '加拿大', '澳大利亚', '新加坡', '荷兰', '瑞典', '芬兰', '挪威',
    '丹麦', '瑞士', '以色列', '巴西', '墨西哥', '阿根廷', '西班牙', '葡萄牙',
    '波兰', '捷克', '匈牙利', '奥地利', '比利时', '土耳其', '沙特', '阿联酋',
    '泰国', '越南', '马来西亚', '印尼', '菲律宾', '南非', '埃及',
    '全球', '国际', '世界', '海外', '亚洲'
]

def extract_doc_year(fname):
    """从文件名提取发布年份，如 '2026全球智数化人才指数报告.pdf' → 2026。"""
    m = re.search(r'(20\d{2})', fname)
    return int(m.group(1)) if m else None

def extract_doc_year_from_content(text, max_scan=800):
    """从文档前若干字符中提取最可能的发布年份（文件名无年份时的兜底）。
    优先识别"XXXX年"格式，取频率最高的年份；完全没有则返回 None。
    """
    sample = text[:max_scan]
    # 优先：明确的"XXXX年"形式
    years_explicit = re.findall(r'(20\d{2})年', sample)
    if years_explicit:
        return int(Counter(years_explicit).most_common(1)[0][0])
    # 次选：裸4位数字（可能是版权年/发布年）
    years_bare = re.findall(r'\b(20\d{2})\b', sample)
    if years_bare:
        return int(Counter(years_bare).most_common(1)[0][0])
    return None

def extract_countries(text):
    """从 chunk 文本中提取提到的国家/地区关键词，返回逗号分隔字符串（最多6个）。"""
    found = []
    for kw in COUNTRY_KEYWORDS:
        if kw in text:
            found.append(kw)
        if len(found) >= 6:
            break
    return ','.join(found) if found else ''

def extract_dominant_year(text):
    """提取 chunk 文本中出现最多的 20xx 年份，无则返回 None。"""
    years = re.findall(r'(20\d{2})', text)
    if not years:
        return None
    return int(Counter(years).most_common(1)[0][0])


# ── ChromaDB REST v2 ──────────────────────────────────

def _col_url(path=""):
    base = f"http://{CHROMA_HOST}:{CHROMA_PORT}/api/v2"
    return f"{base}/tenants/{TENANT}/databases/{DATABASE}/collections{path}"

def _req(method, url, **kwargs):
    r = requests.request(method, url, timeout=15, **kwargs)
    r.raise_for_status()
    return r.json() if r.content else {}

def chroma_heartbeat():
    requests.get(
        f"http://{CHROMA_HOST}:{CHROMA_PORT}/api/v2/heartbeat", timeout=5
    ).raise_for_status()

def chroma_get_or_create():
    try:
        data = _req("GET", _col_url(f"/{COLLECTION}"))
        return data["id"]
    except requests.HTTPError:
        data = _req("POST", _col_url(), json={
            "name": COLLECTION,
            "metadata": {"hnsw:space": "cosine"}
        })
        return data["id"]

def chroma_get_ids(col_id):
    """分页拉取已有知识文档 ID（只拉 k_ 前缀，速度快且不会误判结构化数据 ID）"""
    ids = set()
    offset = 0
    page = 500
    try:
        while True:
            result = _req("POST", _col_url(f"/{col_id}/get"),
                          json={"include": [], "limit": page, "offset": offset})
            batch = result.get("ids", [])
            if not batch:
                break
            # 只记录知识文档的 k_ 前缀 ID，避免结构化数据 ID 干扰跳过逻辑
            ids.update(i for i in batch if i.startswith("k_"))
            offset += len(batch)
            if len(batch) < page:
                break
        return ids
    except Exception:
        return set()

def chroma_update_meta(col_id, ids, metadatas):
    """仅更新 metadata，不修改向量和文档文本。"""
    _req("POST", _col_url(f"/{col_id}/update"), json={
        "ids": ids, "metadatas": metadatas
    })

def chroma_delete_by_filename(col_id, fname):
    """删除指定文件的所有 chunk（用于 --file --force 单文件重建）。
    通过 filename metadata 精准匹配，只删该文件的 chunk，不影响其他文档。
    """
    try:
        _req("POST", _col_url(f"/{col_id}/delete"),
             json={"where": {"filename": {"$eq": fname}}})
        print(f"   🗑️  已删除旧 chunk（filename={fname}）")
    except Exception as e:
        print(f"   ⚠️  删除旧 chunk 失败（将继续入库）: {e}")

def update_metadata_only(col_id):
    """
    --update-meta 模式：从已入库 chunk 的 document 文本和 source 字段重新计算
    doc_year / countries / year_in_chunk，批量写回 metadata，不重新 embedding。
    速度远快于 --force（无 Ollama 调用）。
    """
    print("\n🔄 元数据更新模式（不重新 embedding）...")
    offset, page = 0, 200
    updated, skipped = 0, 0

    while True:
        result = _req("POST", _col_url(f"/{col_id}/get"),
                      json={"where": {"table": {"$eq": "knowledge"}},
                            "include": ["documents", "metadatas"],
                            "limit": page, "offset": offset})
        ids      = result.get("ids", [])
        docs     = result.get("documents", [])
        metas    = result.get("metadatas", [])
        if not ids:
            break

        batch_ids, batch_metas = [], []
        for cid, doc, meta in zip(ids, docs, metas):
            new_meta = dict(meta)   # 保留现有所有字段
            text = str(doc or '')

            # doc_year：优先保留已有值；否则从 source/filename 提取，再从内容推断
            if not new_meta.get("doc_year"):
                fname = new_meta.get("filename", new_meta.get("source", ""))
                yr = extract_doc_year(fname) or extract_doc_year_from_content(text)
                if yr:
                    new_meta["doc_year"] = yr

            # countries
            if not new_meta.get("countries"):
                c = extract_countries(text)
                if c:
                    new_meta["countries"] = c

            # year_in_chunk
            if not new_meta.get("year_in_chunk"):
                y = extract_dominant_year(text)
                if y:
                    new_meta["year_in_chunk"] = y

            if new_meta != meta:
                batch_ids.append(cid)
                batch_metas.append(new_meta)
            else:
                skipped += 1

        if batch_ids:
            chroma_update_meta(col_id, batch_ids, batch_metas)
            updated += len(batch_ids)
            print(f"   ✅ 已更新 {updated} 条...", end="\r")

        offset += len(ids)
        if len(ids) < page:
            break

    print(f"\n   🎉 元数据更新完成：更新 {updated} 条，无需更新 {skipped} 条")

def chroma_delete_knowledge(col_id):
    """删除集合中所有 table=knowledge 的文档（保留结构化数据）"""
    deleted = 0
    page = 200
    while True:
        try:
            result = _req("POST", _col_url(f"/{col_id}/get"),
                          json={"where": {"table": {"$eq": "knowledge"}},
                                "include": [], "limit": page, "offset": 0})
            batch_ids = result.get("ids", [])
            if not batch_ids:
                break
            _req("POST", _col_url(f"/{col_id}/delete"), json={"ids": batch_ids})
            deleted += len(batch_ids)
            print(f"   🗑️  已删除 {deleted} 条...", end="\r")
        except Exception as e:
            print(f"\n   ❌ 删除失败: {e}")
            break
    print(f"\n   ✅ 共删除 {deleted} 条知识文档")
    return deleted

def chroma_add(col_id, ids, embeddings, documents, metadatas):
    _req("POST", _col_url(f"/{col_id}/add"), json={
        "ids": ids, "embeddings": embeddings,
        "documents": documents, "metadatas": metadatas
    })

def chroma_count(col_id):
    return _req("GET", _col_url(f"/{col_id}/count"))


# ── Ollama Embedding ──────────────────────────────────

def _clean_embed_text(text, max_len):
    """清理并截断 embed 文本，避免特殊格式导致 Ollama 500。"""
    t = str(text)[:max_len]
    # 压缩 Markdown 表格分隔行（|---|---| 大量重复导致 tokenizer 异常）
    t = re.sub(r'\|[-:\s|]{4,}', '|', t)
    # 压缩连续重复字符（10个以上压缩为3个）
    t = re.sub(r'(.)\1{9,}', r'\1\1\1', t)
    # 移除控制字符（保留换行）
    t = re.sub(r'[\x00-\x08\x0b-\x1f\x7f]', '', t)
    return t.strip()


def get_embedding(text, retries=3):
    # 失败后逐步截短：4000 → 2500 → 1200，避免特定内容长度触发 Ollama 500
    length_ladder = [4000, 2500, 1200]
    last_exc = None
    for max_len in length_ladder:
        prompt = _clean_embed_text(text, max_len)
        for attempt in range(retries):
            try:
                r = requests.post(
                    f"{OLLAMA_URL}/api/embeddings",
                    json={"model": EMBED_MODEL, "prompt": prompt},
                    timeout=EMBED_TIMEOUT
                )
                r.raise_for_status()
                emb = r.json().get("embedding")
                if not emb:
                    raise ValueError("embedding 返回为空")
                if max_len < 4000:
                    print(f"  ✅ 截短至 {max_len} 字后成功")
                return emb
            except Exception as e:
                last_exc = e
                if attempt < retries - 1:
                    is_500 = hasattr(e, 'response') and getattr(e.response, 'status_code', 0) == 500
                    wait = (attempt + 1) * 5 if is_500 else 2 ** attempt
                    print(f"  ⚠️  embedding 失败（{max_len}字，第{attempt+1}次），{wait}s 后重试: {e}")
                    time.sleep(wait)
                    if is_500 and attempt == 0:
                        refresh_embed_model()
        # 当前长度全部失败，缩短后重试
        print(f"  ⚠️  {max_len}字长度全失败，尝试截短...")
        refresh_embed_model()
    raise last_exc


# ── 文本提取 ──────────────────────────────────────────

def clean_page_text(text):
    """清理 PDF 文本常见噪声"""
    if not text:
        return ''
    # 修复连字符换行（英文/中文均可能出现）
    text = re.sub(r'([A-Za-z一-龥])-\n([A-Za-z一-龥])', r'\1\2', text)
    # 删除独占一行的孤立页码数字（页脚/页眉）
    text = re.sub(r'(?m)^\s*\d{1,3}\s*$', '', text)
    # 压缩多余空行
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def table_to_markdown(table_data):
    """把 pdfplumber 提取的原始表格转成 Markdown，多级表头自动压平。
    返回 Markdown 字符串，或 None（表格无实质内容时）。
    """
    if not table_data:
        return None

    def clean_cell(c):
        return re.sub(r'\s+', ' ', str(c or '').strip())

    rows = [[clean_cell(c) for c in row] for row in table_data
            if any(str(c or '').strip() for c in row)]
    if len(rows) < 2:
        return None

    # 判断是否为表头行（数字占比 < 40% → 视为表头）
    def is_header_row(row):
        non_empty = [c for c in row if c]
        if not non_empty:
            return True
        numeric = sum(1 for c in non_empty
                      if re.match(r'^[\d.,\-%+()（）年月日]+$', c))
        return numeric / len(non_empty) < 0.4

    header_rows, data_start = [], 0
    for i, row in enumerate(rows):
        if i < 5 and is_header_row(row):
            header_rows.append(row)
            data_start = i + 1
        else:
            break
    if not header_rows:
        header_rows, data_start = [rows[0]], 1

    col_count = max(len(r) for r in header_rows)

    # 向右传播合并单元格的值
    def propagate(row, n):
        row = list(row) + [''] * max(0, n - len(row))
        last = ''
        for i, c in enumerate(row):
            if c:
                last = c
            else:
                row[i] = last
        return row

    expanded = [propagate(r, col_count) for r in header_rows]

    headers = []
    for col in range(col_count):
        parts = []
        for row in expanded:
            val = row[col] if col < len(row) else ''
            if val and (not parts or val != parts[-1]):
                parts.append(val)
        headers.append('/'.join(parts) if parts else f'列{col + 1}')

    data_rows = rows[data_start:]
    if not data_rows:
        return None

    lines = ['| ' + ' | '.join(headers) + ' |',
             '| ' + ' | '.join(['---'] * len(headers)) + ' |']
    for row in data_rows:
        cells = (list(row) + [''] * len(headers))[:len(headers)]
        if any(cells):
            lines.append('| ' + ' | '.join(cells) + ' |')

    return '\n'.join(lines) if len(lines) > 2 else None


def heading_level(line):
    """返回标题级别：0=非标题，1=一级，2=二级，3=三级。
    H1：第X章/节，或汉字数字+顿号（一、二、三、四、）
    H2：（一）（二）... 括号包裹的编号
    H3：1. / 1.1 / 3.2. 等纯数字编号（后接中文）
    图表标题不计入层级（返回0，不影响标题栈）
    """
    line = line.strip()
    if not line or len(line) > 60 or len(line) < 2:
        return 0
    if line[-1] in '。，；：、.,;）)】%|':
        return 0
    if re.match(r'^[%\-\|]', line):
        return 0
    if re.match(r'^(图表|图|表)\s*\d+', line):
        return 0   # 图表标题不算章节
    if re.match(r'^(第[一二三四五六七八九十百\d]+[章节部分篇]|[一二三四五六七八九十]+[、．\.][\s一-鿿])', line):
        return 1
    if re.match(r'^[（(【]\s*[一二三四五六七八九十\d]+\s*[)）】][\s一-鿿]', line):
        return 2
    if re.match(r'^\d+(\.\d+)*[\.。\s、]+[一-鿿]', line):
        return 3
    return 0

def is_heading_line(line):
    """向后兼容：只需知道是否为标题"""
    return heading_level(line) > 0


def detect_printed_page(page_text, physical_num):
    """从页面文本中识别印刷页码（页脚/页眉的孤立数字）。
    找不到时回退到物理页码。
    """
    lines = [l.strip() for l in page_text.split('\n') if l.strip()]
    # 优先检查最后5行（页脚）
    for line in reversed(lines[-5:] if len(lines) >= 5 else lines):
        if re.match(r'^\d{1,3}$', line):
            n = int(line)
            if 1 <= n <= physical_num + 100:
                return n
    # 再检查前3行（页眉）
    for line in lines[:3]:
        if re.match(r'^\d{1,3}$', line):
            n = int(line)
            if 1 <= n <= physical_num + 100:
                return n
    return physical_num


def render_page_image_b64(path, phys_num, scale=2):
    """用 pymupdf 把 PDF 某物理页渲染成 PNG，返回 base64 字符串。"""
    try:
        import fitz
        doc = fitz.open(path)
        page = doc[phys_num - 1]
        pix = page.get_pixmap(matrix=fitz.Matrix(scale, scale))
        b64 = base64.b64encode(pix.tobytes('png')).decode()
        doc.close()
        return b64
    except Exception as e:
        print(f'  ⚠️  页面渲染失败: {e}')
        return None


def describe_with_llava(img_b64, title_hint=''):
    """调用本地 Ollama llava 模型，把图表页转成结构化文字描述。"""
    prompt = (
        '请详细分析这张图表，用中文输出以下内容：\n'
        '1. 图表类型和标题\n'
        '2. 所有数据数值（格式：分类：数值，逐一列出，不要遗漏）\n'
        '3. 主要结论或趋势\n'
        + (f'图表标题参考：{title_hint}\n' if title_hint else '')
        + '请精确读取每个数字。'
    )
    try:
        r = requests.post(
            f'{OLLAMA_URL}/api/generate',
            json={'model': 'llava', 'prompt': prompt,
                  'images': [img_b64], 'stream': False},
            timeout=120
        )
        r.raise_for_status()
        response = r.json().get('response', '').strip()
        return f'[图表内容]\n{response}' if response else ''
    except requests.exceptions.ConnectionError:
        return ''  # llava 未安装，静默跳过
    except Exception as e:
        print(f'  ⚠️  llava 识别失败: {e}')
        return ''


def unload_llava():
    """卸载 llava，释放内存给 embedding 模型。在同一 PDF 所有图表页处理完后调用一次。"""
    try:
        requests.post(f'{OLLAMA_URL}/api/generate',
                      json={'model': 'llava', 'keep_alive': 0}, timeout=10)
        time.sleep(3)   # 等待 Ollama 实际释放 VRAM/RAM
    except Exception:
        pass


def refresh_embed_model():
    """主动卸载 nomic-embed-text，防止长时间运行后内存碎片积累导致 500 OOM。
    每处理约 80 个 chunk 后调用一次；下次 get_embedding 会自动重新加载模型。"""
    try:
        requests.post(f'{OLLAMA_URL}/api/generate',
                      json={'model': EMBED_MODEL, 'keep_alive': 0}, timeout=10)
        time.sleep(2)
    except Exception:
        pass


def extract_pdf(path):
    """提取 PDF 全文（兼容旧调用）"""
    pages = extract_pdf_pages(path)
    return "\n".join(text for _, text, _ in pages)


def extract_pdf_pages(path):
    """按页提取 PDF，返回 [(printed_page_num, text, section_heading), ...].

    处理：
      - 表格 → Markdown（多级表头压平，避免重复提取）
      - 正文去掉表格区域（防止内容重复）
      - 图片页：文字极少时调 llava 识别
      - 印刷页码识别（复用原始文本，不重复调用）
      - 章节标题追踪（跨页维持状态，供 chunk 上下文使用）
      - 文字清洗（连字符换行、孤立页码、多余空行）
    优先 pdfplumber；回退 pymupdf。
    """
    # ── pdfplumber ────────────────────────────────────────
    try:
        import pdfplumber
        pages = []
        heading_stack = ['', '', '']   # [h1, h2, h3]，跨页维持多级标题状态

        def update_heading_stack(line):
            lv = heading_level(line)
            if lv == 1:
                heading_stack[0] = line.strip()
                heading_stack[1] = ''
                heading_stack[2] = ''
            elif lv == 2:
                heading_stack[1] = line.strip()
                heading_stack[2] = ''
            elif lv == 3:
                heading_stack[2] = line.strip()

        def current_heading_path():
            parts = [h for h in heading_stack if h]
            return ' > '.join(parts)

        with pdfplumber.open(path) as pdf:
            for phys_num, page in enumerate(pdf.pages, start=1):
                parts = []

                # 0. 先提取完整原始文本（供页码检测用，不重复调用）
                raw_text = page.extract_text() or ''

                # 1. 找出所有表格及其 bbox
                found_tables = page.find_tables()
                table_bboxes = [t.bbox for t in found_tables]

                # 2. 每张表格转 Markdown
                for t in found_tables:
                    md = table_to_markdown(t.extract())
                    if md:
                        parts.append(md)

                # 3. 提取表格区域以外的正文（避免重复）
                if table_bboxes:
                    bboxes = table_bboxes   # 闭包捕获当前页的 bbox 列表
                    def not_in_table(obj):
                        for (x0, top, x1, bottom) in bboxes:
                            if (obj.get('x0', 0) >= x0 - 2 and
                                    obj.get('x1', 0) <= x1 + 2 and
                                    obj.get('top', 0) >= top - 2 and
                                    obj.get('bottom', 0) <= bottom + 2):
                                return False
                        return True
                    body_text = page.filter(not_in_table).extract_text() or ''
                else:
                    body_text = raw_text

                body_text = clean_page_text(body_text)

                # 4. 从正文中检测章节标题，跨页维持多级标题栈
                for line in body_text.split('\n'):
                    update_heading_stack(line)

                if body_text:
                    parts.append(body_text)

                # 5. 图片页检测：文字极少且有图 → 调 llava 识别
                img_count = len(page.images)
                text_too_short = len(body_text) < 80 and not found_tables
                if img_count and text_too_short and not use_no_vision:
                    print(f'\n   🖼️  第{phys_num}页图表页，调用 llava...', end='', flush=True)
                    img_b64 = render_page_image_b64(path, phys_num)
                    if img_b64:
                        desc = describe_with_llava(img_b64, body_text[:120])
                        if desc:
                            parts.append(desc)
                            print(' ✅')
                        else:
                            parts.append(f'[本页含图表（共 {img_count} 张），llava 未安装或识别失败]')
                            print(' ⚠️  llava 无输出（是否已 ollama pull llava？）')
                    else:
                        parts.append(f'[本页含图表（共 {img_count} 张），页面渲染失败]')

                combined = '\n\n'.join(p for p in parts if p.strip())
                if not combined:
                    continue

                # 6. 复用 raw_text 识别印刷页码（不再重复调用 extract_text）
                printed = detect_printed_page(raw_text, phys_num)
                pages.append((printed, combined, current_heading_path()))

        return pages

    except ImportError:
        pass  # 回退 pymupdf
    except Exception as e:
        print(f"  ⚠️  pdfplumber 失败，尝试 pymupdf: {e}")

    # ── pymupdf 回退 ──────────────────────────────────────
    try:
        import fitz
        doc = fitz.open(path)
        pages = []
        fb_stack = ['', '', '']
        def fb_update(line):
            lv = heading_level(line)
            if lv == 1: fb_stack[0] = line.strip(); fb_stack[1] = ''; fb_stack[2] = ''
            elif lv == 2: fb_stack[1] = line.strip(); fb_stack[2] = ''
            elif lv == 3: fb_stack[2] = line.strip()
        for phys_num, page in enumerate(doc, start=1):
            raw_text = page.get_text('text')
            text = clean_page_text(raw_text)
            if not text:
                continue
            for line in text.split('\n'):
                fb_update(line)
            printed = detect_printed_page(raw_text, phys_num)
            pages.append((printed, text, ' > '.join(h for h in fb_stack if h)))
        doc.close()
        return pages
    except ImportError:
        print('  ⚠️  未安装 pdfplumber 和 pymupdf，跳过此文件')
        return []
    except Exception as e:
        print(f'  ❌ PDF 提取失败: {e}')
        return []

def extract_docx(path):
    """提取 Word 文档，追踪章节标题，返回 [(chunk_text, page_num, section_heading), ...]。
    page_num 始终为 None（Word 无页码概念）；section 随段落样式推进。
    """
    try:
        from docx import Document
        doc = Document(path)
        dx_stack = ['', '', '']   # [h1, h2, h3]
        HEADING_STYLES_BY_LEVEL = {
            'heading 1': 1, 'heading 2': 2, 'heading 3': 3,
            '标题 1': 1, '标题 2': 2, '标题 3': 3,
            '标题1': 1, '标题2': 2, '标题3': 3,
        }
        def docx_update_stack(text, style_name):
            lv = HEADING_STYLES_BY_LEVEL.get(style_name, 0) or heading_level(text)
            if lv == 1: dx_stack[0] = text; dx_stack[1] = ''; dx_stack[2] = ''
            elif lv == 2: dx_stack[1] = text; dx_stack[2] = ''
            elif lv == 3: dx_stack[2] = text
        def docx_heading_path():
            return ' > '.join(h for h in dx_stack if h)

        segments = []   # [(text, heading_path)]
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style_name = (para.style.name or '').lower()
            docx_update_stack(text, style_name)
            segments.append((text, docx_heading_path()))

        # 追加表格
        for table in doc.tables:
            rows = [[c.text.strip() for c in row.cells] for row in table.rows]
            rows = [r for r in rows if any(r)]
            md = table_to_markdown(rows)
            if md:
                segments.append((md, docx_heading_path()))

        # 把相邻同章节段落合并后再分块，保持章节上下文
        result = []
        group_text, group_heading = [], ''
        for text, heading in segments:
            if heading != group_heading and group_text:
                for c in chunk_text('\n'.join(group_text)):
                    result.append((c, None, group_heading))
                group_text = []
            group_heading = heading
            group_text.append(text)
        if group_text:
            for c in chunk_text('\n'.join(group_text)):
                result.append((c, None, group_heading))

        return result
    except Exception as e:
        print(f"  ❌ Word 提取失败: {e}")
        return []

def extract_xlsx(path):
    """把 Excel 每行转成自然语言句子，跳过图表类型汇总等非人才数据文件"""
    fname = os.path.basename(path)
    # 跳过：图表类型汇总（开发参考）和基础数据汇总（已由 ingest.js 处理）
    if '图表类型' in fname or '基础数据汇总' in fname:
        return ""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        parts = []
        for ws in wb.worksheets:
            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                continue
            # 第一行判断是否为表头
            header = [str(c).strip() if c is not None else '' for c in rows[0]]
            has_header = any(h for h in header)
            data_rows = rows[1:] if has_header else rows
            for row in data_rows:
                cells = [str(c).strip() if c is not None else '' for c in row]
                if not any(cells):
                    continue
                if has_header:
                    # 拼成"字段名：值"的自然语言句子
                    pairs = [f"{h}：{v}" for h, v in zip(header, cells) if h and v and v != 'None']
                    if pairs:
                        parts.append('，'.join(pairs) + '。')
                else:
                    parts.append('，'.join(c for c in cells if c and c != 'None') + '。')
        wb.close()
        return "\n".join(parts)
    except ImportError:
        print("  ⚠️  未安装 openpyxl，跳过此文件")
        return ""
    except Exception as e:
        print(f"  ❌ Excel 提取失败: {e}")
        return ""


# ── 分块 ──────────────────────────────────────────────

def preprocess_talent_list(text):
    """检测名单类密集文本，按年份自动断行"""
    year_count = len(re.findall(r'20\d{2}年', text))
    if year_count < 3 or len(text) / max(year_count, 1) > 80:
        return text
    text = re.sub(r'\s+', '', text)
    text = re.sub(r'(20\d{2}年(?:度|入选|获|荣获|被评|当选))', r'\n\1', text)
    text = re.sub(r'([^\n])(20\d{2}年[^度入获荣被当\d])', r'\1\n\2', text)
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    return '\n'.join(lines)

def _chunk_plain(text):
    """对纯文本做滑动窗口分块（内部用）"""
    text = text.strip()
    if not text:
        return []
    if len(text) <= CHUNK_SIZE:
        return [text]
    lines = text.split('\n')
    if len(lines) > 10:
        chunks, current, cur_len = [], [], 0
        for line in lines:
            if cur_len + len(line) > CHUNK_SIZE and current:
                chunks.append('\n'.join(current))
                overlap, ol = [], 0
                for l in reversed(current):
                    if ol + len(l) > CHUNK_OVERLAP:
                        break
                    overlap.insert(0, l)
                    ol += len(l)
                current, cur_len = overlap, ol
            current.append(line)
            cur_len += len(line)
        if current:
            chunks.append('\n'.join(current))
        return [c for c in chunks if len(c.strip()) > 20]
    chunks, start = [], 0
    while start < len(text):
        chunk = text[start:start + CHUNK_SIZE].strip()
        if chunk:
            chunks.append(chunk)
        start += CHUNK_SIZE - CHUNK_OVERLAP
    return chunks


def chunk_text(text):
    """分块，Markdown 表格保持整体不切割（切断后列头丢失无意义）。
    大表格（>CHUNK_SIZE）按行分组，但每组都携带表头。
    """
    text = preprocess_talent_list(text)
    text = re.sub(r'\n{3,}', '\n\n', text).strip()
    if not text:
        return []

    # 把内容拆成"表格段"和"文本段"交替序列
    segments = []          # [(type, lines)]
    cur_type, cur_lines = None, []
    for line in text.split('\n'):
        ltype = 'table' if line.startswith('|') else 'text'
        if ltype != cur_type and cur_lines:
            segments.append((cur_type, cur_lines))
            cur_lines = []
        cur_type = ltype
        cur_lines.append(line)
    if cur_lines:
        segments.append((cur_type, cur_lines))

    chunks = []
    for seg_type, seg_lines in segments:
        seg_text = '\n'.join(seg_lines).strip()
        if not seg_text:
            continue

        if seg_type == 'text':
            chunks.extend(_chunk_plain(seg_text))
        else:
            # Markdown 表格：提取表头行（前两行：列名 + 分隔线）
            header_lines = seg_lines[:2]
            data_lines   = seg_lines[2:]

            if len(seg_text) <= CHUNK_SIZE or not data_lines:
                # 整张表放一个 chunk
                chunks.append(seg_text)
            else:
                # 按行分组，每组带表头，避免超长
                group, g_len = list(header_lines), sum(len(l) for l in header_lines)
                for dline in data_lines:
                    if g_len + len(dline) > CHUNK_SIZE and len(group) > 2:
                        chunks.append('\n'.join(group))
                        group = list(header_lines)   # 新 chunk 重新带表头
                        g_len = sum(len(l) for l in header_lines)
                    group.append(dline)
                    g_len += len(dline)
                if len(group) > 2:
                    chunks.append('\n'.join(group))

    return [c for c in chunks if len(c.strip()) > 20]


# ── 文件扫描 ──────────────────────────────────────────

def scan_docs():
    docs = []
    for root, dirs, files in os.walk(DOCS_FOLDER):
        dirs[:] = [d for d in dirs if not d.startswith('.')]
        for fname in sorted(files):
            if fname.startswith('~$') or fname.startswith('.'):
                continue
            if os.path.splitext(fname)[1].lower() in ('.pdf', '.docx', '.xlsx'):
                docs.append(os.path.join(root, fname))
    return docs

def make_id(path, chunk_text_content):
    """用文件路径 + chunk 内容哈希生成稳定 ID。
    内容不变则 ID 不变，文件改动只影响变化的 chunk。
    """
    rel = os.path.relpath(path, DOCS_FOLDER)
    content_hash = hashlib.md5(f"{rel}|{chunk_text_content}".encode()).hexdigest()[:20]
    return "k_" + content_hash


# ── 主流程 ────────────────────────────────────────────

def main():
    import sys
    force         = "--force"     in sys.argv
    update_meta   = "--update-meta" in sys.argv
    global use_no_vision
    use_no_vision = "--no-vision" in sys.argv

    # --file <文件名> 只处理指定文件（支持部分名匹配）
    file_filter = None
    if "--file" in sys.argv:
        idx = sys.argv.index("--file")
        if idx + 1 < len(sys.argv):
            file_filter = sys.argv[idx + 1].strip()

    print("🚀 知识库文档入库")
    print(f"   文档目录 : {DOCS_FOLDER}")
    print(f"   ChromaDB : {CHROMA_HOST}:{CHROMA_PORT}")
    print(f"   Ollama   : {OLLAMA_URL}  [{EMBED_MODEL}]")
    print(f"   图表识别 : {'关闭（--no-vision）' if use_no_vision else '开启（llava，未安装则自动跳过）'}")
    if force:
        print("   ⚠️  --force 模式：将删除旧知识文档后重新入库")
    if file_filter:
        print(f"   🔍 单文件模式：{file_filter}")

    all_files = scan_docs()
    files = [f for f in all_files if file_filter is None or file_filter in os.path.basename(f)] if file_filter else all_files
    if not files:
        print(f"\n❌ 未找到匹配文档（filter={file_filter}）")
        return
    print(f"\n📄 找到 {len(files)} 个文档{f'（共 {len(all_files)} 个，已过滤）' if file_filter else ''}")

    print("\n🔗 连接 ChromaDB...")
    try:
        chroma_heartbeat()
        col_id = chroma_get_or_create()
        if update_meta:
            update_metadata_only(col_id)
            print(f"   ChromaDB 总量: {chroma_count(col_id)} 条")
            return
        if force and file_filter:
            # 单文件重建：只删该文件的 chunk，不影响其他文档
            for f in files:
                chroma_delete_by_filename(col_id, os.path.basename(f))
        elif force:
            print("   🗑️  删除旧知识文档...")
            chroma_delete_knowledge(col_id)
        # 只拉取 k_ 前缀 ID（知识文档），避免结构化数据 ID 干扰跳过逻辑
        existing = chroma_get_ids(col_id)
        print(f"   ✅ 已有知识文档 ID: {len(existing)} 条")
    except Exception as e:
        print(f"   ❌ 连接失败: {e}")
        return

    print("\n🧪 测试 Ollama...")
    try:
        dim = len(get_embedding("测试"))
        print(f"   ✅ embedding 维度: {dim}")
    except Exception as e:
        print(f"   ❌ Ollama 失败: {e}")
        return

    total_new, skipped, failed = 0, 0, []

    for path in files:
        rel  = os.path.relpath(path, DOCS_FOLDER)
        ext  = os.path.splitext(path)[1].lower()
        fname = os.path.basename(path)
        print(f"\n📖 {rel}")

        # ── 提取文本，PDF 保留页码和章节信息 ────────────
        if ext == '.pdf':
            # [(page_num, text, section_heading), ...]
            pdf_pages = extract_pdf_pages(path)
            if not pdf_pages:
                print("   ⚠️  内容过短，跳过")
                failed.append(rel)
                continue
            # [(chunk_text, page_num, section_heading), ...]
            chunks_with_page = []
            for page_num, page_text, heading in pdf_pages:
                for c in chunk_text(page_text):
                    chunks_with_page.append((c, page_num, heading))
            total_chars = sum(len(t) for _, t, _ in pdf_pages)
            print(f"   {total_chars} 字 → {len(chunks_with_page)} 块（共 {len(pdf_pages)} 页）")
            # 所有图表页处理完后统一卸载 llava，腾出内存给 embedding
            unload_llava()
        else:
            if ext == '.xlsx':
                text = extract_xlsx(path)
                if len(text.strip()) < 50:
                    print("   ⚠️  内容过短，跳过")
                    failed.append(rel)
                    continue
                chunks_with_page = [(c, None, '') for c in chunk_text(text)]
                print(f"   {len(text)} 字 → {len(chunks_with_page)} 块")
            else:  # .docx
                chunks_with_page = extract_docx(path)
                if not chunks_with_page:
                    print("   ⚠️  内容过短，跳过")
                    failed.append(rel)
                    continue
                total_chars = sum(len(c) for c, _, _ in chunks_with_page)
                print(f"   ~{total_chars} 字 → {len(chunks_with_page)} 块")

        ids, embs, docs, metas = [], [], [], []
        new_count = 0
        total_chunks = len(chunks_with_page)
        doc_prefix = f"【来源：{fname}】\n"   # 文档级前缀，只加在 embedding 里

        # 发布年份：优先从文件名提取，兜底从文档内容第一段提取
        doc_year = extract_doc_year(fname)
        if doc_year is None and chunks_with_page:
            first_text = ' '.join(c for c, _, _ in chunks_with_page[:3])
            doc_year = extract_doc_year_from_content(first_text)
            if doc_year:
                print(f"   📅 年份从内容推断：{doc_year}")

        embed_count = 0  # 记录本文件实际 embedding 次数，用于定期刷新模型
        for i, (chunk, page_num, section_heading) in enumerate(chunks_with_page):
            cid = make_id(path, chunk)
            if cid in existing:
                skipped += 1
                continue

            # 每 80 次 embedding 主动刷新模型，防止内存碎片积累导致 500 OOM
            if embed_count > 0 and embed_count % 80 == 0:
                print(f"   🔄 定期刷新 embedding 模型（已嵌入 {embed_count} 块）...")
                refresh_embed_model()

            # 存储文本：章节标题 + chunk（LLM 看到上下文）
            section_prefix = f"【章节：{section_heading}】\n" if section_heading else ''
            stored_text = section_prefix + chunk

            # 嵌入文本：文件名 + 章节 + chunk（向量同时捕获文档和章节归属）
            embed_text = doc_prefix + stored_text

            try:
                emb = get_embedding(embed_text)
                embed_count += 1
            except Exception as e:
                print(f"   ❌ embedding chunk {i}: {e}")
                time.sleep(1)
                continue

            ids.append(cid); embs.append(emb); docs.append(stored_text)
            meta = {"table": "knowledge", "source": rel, "filename": fname, "chunk": i}
            if doc_year is not None:
                meta["doc_year"] = doc_year
            if page_num is not None:
                meta["page"] = page_num
            if section_heading:
                meta["section"] = section_heading
            chunk_countries = extract_countries(chunk)
            if chunk_countries:
                meta["countries"] = chunk_countries
            chunk_year = extract_dominant_year(chunk)
            if chunk_year is not None:
                meta["year_in_chunk"] = chunk_year
            metas.append(meta)
            new_count += 1

            if len(ids) >= BATCH_SIZE:
                try:
                    chroma_add(col_id, ids, embs, docs, metas)
                    existing.update(ids)
                    print(f"   ✅ {new_count}/{total_chunks}", end="\r")
                except Exception as e:
                    print(f"\n   ❌ 写入失败: {e}")
                    # 写入失败不清空，下次循环继续积累，避免数据丢失
                    continue
                ids, embs, docs, metas = [], [], [], []

        if ids:
            try:
                chroma_add(col_id, ids, embs, docs, metas)
                existing.update(ids)
            except Exception as e:
                print(f"\n   ❌ 写入尾批失败（{len(ids)} 条）: {e}")

        total_new += new_count
        print(f"   ✅ 新增 {new_count} 块")

    print(f"\n{'='*50}")
    print(f"🎉 完成  新增 {total_new} 块，跳过 {skipped} 块")
    if failed:
        print(f"   ⚠️  提取失败: {failed}")
    print(f"   ChromaDB 总量: {chroma_count(col_id)} 条")


if __name__ == "__main__":
    main()
