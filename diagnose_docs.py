"""
diagnose_docs.py — 快速诊断 资料/ 下每个文件的文本提取情况
不需要 ChromaDB 或 Ollama，直接运行即可

用法：python diagnose_docs.py
"""

import os
import sys

DOCS_FOLDER = os.path.join(os.path.dirname(__file__), "public", "资料")

def extract_pdf(path):
    try:
        import fitz
        doc = fitz.open(path)
        pages_with_text = 0
        pages_empty = 0
        total_text = []
        for i, page in enumerate(doc):
            t = page.get_text("text").strip()
            if t:
                pages_with_text += 1
                total_text.append(t)
            else:
                pages_empty += 1
        doc.close()
        text = "\n".join(total_text)
        return text, pages_with_text, pages_empty
    except ImportError:
        return "", 0, 0
    except Exception as e:
        return f"[ERROR: {e}]", 0, 0

def extract_docx(path):
    try:
        from docx import Document
        doc = Document(path)
        parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        text = "\n".join(parts)
        return text, len(parts), 0
    except Exception as e:
        return f"[ERROR: {e}]", 0, 0

def extract_xlsx(path):
    fname = os.path.basename(path)
    if '图表类型' in fname or '基础数据汇总' in fname:
        return "[跳过（在跳过列表中）]", 0, 0
    try:
        import openpyxl
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        rows_count = sum(1 for ws in wb.worksheets for _ in ws.iter_rows())
        wb.close()
        return f"[Excel，{rows_count}行]", rows_count, 0
    except Exception as e:
        return f"[ERROR: {e}]", 0, 0

def chunk_count_estimate(text_len, chunk_size=600, overlap=100):
    if text_len <= chunk_size:
        return 1
    step = chunk_size - overlap
    return max(1, (text_len - chunk_size) // step + 1)

def main():
    print(f"📂 扫描目录: {DOCS_FOLDER}\n")

    files = []
    for root, dirs, fnames in os.walk(DOCS_FOLDER):
        dirs[:] = [d for d in dirs if not d.startswith('.')]
        for fname in sorted(fnames):
            if fname.startswith('~$') or fname.startswith('.'):
                continue
            ext = os.path.splitext(fname)[1].lower()
            if ext in ('.pdf', '.docx', '.xlsx'):
                files.append(os.path.join(root, fname))

    print(f"找到 {len(files)} 个文件\n")
    print("=" * 70)

    total_chunks = 0
    for path in files:
        rel = os.path.relpath(path, DOCS_FOLDER)
        ext = os.path.splitext(path)[1].lower()

        if ext == '.pdf':
            text, pages_ok, pages_empty = extract_pdf(path)
            pages_info = f"文字页:{pages_ok} / 空白页:{pages_empty}"
        elif ext == '.docx':
            text, pages_ok, pages_empty = extract_docx(path)
            pages_info = f"段落:{pages_ok}"
        else:
            text, pages_ok, pages_empty = extract_xlsx(path)
            pages_info = ""

        if text.startswith("["):
            # Error or skip message
            print(f"⚠️  {rel}")
            print(f"   {text}")
        else:
            char_count = len(text.strip())
            est_chunks = chunk_count_estimate(char_count)
            total_chunks += est_chunks
            status = "✅" if char_count >= 50 else "❌ 内容过短(<50字)"
            print(f"{status}  {rel}")
            print(f"   {pages_info}  |  字符数: {char_count:,}  |  预计块数: {est_chunks}")
            if char_count >= 50:
                # 显示前100字预览
                preview = text.strip()[:150].replace('\n', ' ')
                print(f"   预览: {preview}...")
                # 搜索德国相关内容
                if '德国' in text or 'Germany' in text or 'German' in text:
                    idx = text.find('德国') if '德国' in text else text.find('Germany')
                    snippet = text[max(0, idx-20):idx+100].replace('\n', ' ')
                    print(f"   🔍 含「德国」: ...{snippet}...")
        print()

    print("=" * 70)
    print(f"📊 预计总块数: {total_chunks}")
    print("\n提示：预计块数 >> 实际入库块数 → 说明入库过程有写入失败或大量跳过")

if __name__ == "__main__":
    main()
